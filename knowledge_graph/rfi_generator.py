#!/usr/bin/env python3
"""
RFI Response Generator Agent
Automatically drafts professional responses to Request for Information (RFI)

Features:
- Extracts RFI questions from opportunity documents
- Drafts responses using Claude AI
- References company capabilities and past performance
- Generates properly formatted Word documents
- Includes compliance matrices
"""

import sys
import os
import re
from typing import Dict, List
from datetime import datetime
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# For Claude API
try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False
    print("⚠️  anthropic package not installed. Install with: pip install anthropic")

# For Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not installed. Install with: pip install python-docx")


class CompanyProfile:
    """Company information for RFI responses"""
    
    def __init__(self):
        self.company_name = "Your Company Name"
        self.duns = "123456789"
        self.cage_code = "1A2B3"
        self.address = "123 Main Street, City, ST 12345"
        self.phone = "(555) 123-4567"
        self.website = "www.yourcompany.com"
        self.socioeconomic_status = "Small Business"  # or 8(a), SDVOSB, etc.
        
        # Core capabilities
        self.capabilities = [
            "Cloud computing and migration services",
            "Cybersecurity assessment and implementation",
            "Software development and integration",
            "Data analytics and AI/ML solutions",
            "IT infrastructure modernization",
            "Program management and support"
        ]
        
        # Past performance (in production, load from database)
        self.past_performance = [
            {
                'project': 'DOD Cloud Migration',
                'client': 'Department of Defense',
                'value': '$2.5M',
                'period': '2022-2024',
                'description': 'Migrated legacy systems to AWS cloud infrastructure',
                'poc': 'John Smith, john.smith@dod.mil, (555) 111-2222'
            },
            {
                'project': 'NASA IT Modernization',
                'client': 'NASA',
                'value': '$1.8M',
                'period': '2021-2023',
                'description': 'Modernized enterprise applications and infrastructure',
                'poc': 'Jane Doe, jane.doe@nasa.gov, (555) 333-4444'
            },
            {
                'project': 'DHS Cybersecurity Assessment',
                'client': 'Department of Homeland Security',
                'value': '$950K',
                'period': '2023-2024',
                'description': 'Security assessment and penetration testing services',
                'poc': 'Bob Johnson, bob.johnson@dhs.gov, (555) 555-6666'
            }
        ]
        
        # Staff/Team
        self.key_personnel = [
            {
                'name': 'John Smith',
                'title': 'Program Manager',
                'clearance': 'Secret',
                'experience': '15 years in federal IT programs'
            },
            {
                'name': 'Sarah Johnson',
                'title': 'Technical Lead',
                'clearance': 'Top Secret',
                'experience': '12 years in cloud architecture and cybersecurity'
            }
        ]
        
        # Certifications
        self.company_certifications = [
            'ISO 9001:2015',
            'ISO 27001:2013',
            'CMMI Level 3',
            'SAM.gov Registered'
        ]


class RFIResponseGenerator:
    """Generates professional RFI responses"""
    
    def __init__(self):
        self.company = CompanyProfile()
        
        # Initialize Claude
        if ANTHROPIC_AVAILABLE:
            api_key = os.getenv('ANTHROPIC_API_KEY')
            if api_key:
                self.client = anthropic.Anthropic(api_key=api_key)
                self.claude_available = True
            else:
                print("⚠️  ANTHROPIC_API_KEY not found in .env")
                self.claude_available = False
        else:
            self.claude_available = False
    
    def extract_rfi_questions(self, opportunity: Dict) -> List[Dict]:
        """Extract RFI questions from opportunity description"""
        
        description = opportunity.get('description', '')
        
        # Common RFI question patterns
        questions = []
        
        # Look for numbered questions
        numbered_pattern = r'(\d+)\.\s*(.+?)(?=\d+\.|$)'
        matches = re.findall(numbered_pattern, description, re.DOTALL)
        
        for num, question in matches:
            questions.append({
                'number': num,
                'question': question.strip(),
                'type': 'general'
            })
        
        # Common RFI categories if no numbered questions found
        if not questions:
            # Create default RFI sections
            questions = [
                {
                    'number': '1',
                    'question': 'Company Background and Qualifications',
                    'type': 'company_info'
                },
                {
                    'number': '2',
                    'question': 'Technical Approach and Capabilities',
                    'type': 'technical'
                },
                {
                    'number': '3',
                    'question': 'Past Performance and References',
                    'type': 'past_performance'
                },
                {
                    'number': '4',
                    'question': 'Key Personnel',
                    'type': 'personnel'
                }
            ]
        
        return questions
    
    def generate_response_with_claude(self, question: Dict, opportunity: Dict) -> str:
        """Generate response using Claude AI"""
        
        if not self.claude_available:
            return "[Claude AI not available - placeholder response]"
        
        # Build context for Claude
        company_context = f"""
Company: {self.company.company_name}
Capabilities: {', '.join(self.company.capabilities)}
Socioeconomic Status: {self.company.socioeconomic_status}
Certifications: {', '.join(self.company.company_certifications)}
        """
        
        past_perf_context = "\n".join([
            f"- {pp['project']} for {pp['client']} ({pp['period']}): {pp['description']}"
            for pp in self.company.past_performance
        ])
        
        prompt = f"""You are writing a professional response to a Request for Information (RFI) for a federal government contract opportunity.

Opportunity: {opportunity.get('title', 'Unknown')}
Agency: {opportunity.get('agency', 'Unknown')}

Our Company Profile:
{company_context}

Our Past Performance:
{past_perf_context}

RFI Question:
{question['question']}

Please write a professional, concise response (2-3 paragraphs) that:
1. Directly answers the question
2. Highlights relevant company capabilities
3. References specific past performance when applicable
4. Uses professional federal contracting language
5. Is confident but not boastful

Write the response now:"""
        
        try:
            message = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=1000,
                messages=[{"role": "user", "content": prompt}]
            )
            
            response_text = message.content[0].text
            return response_text
            
        except Exception as e:
            print(f"Error calling Claude API: {e}")
            return self.generate_fallback_response(question)
    
    def generate_fallback_response(self, question: Dict) -> str:
        """Generate a basic response without AI"""
        
        q_type = question.get('type', 'general')
        
        if q_type == 'company_info':
            return f"""{self.company.company_name} is a {self.company.socioeconomic_status} providing federal IT services. 
We specialize in {', '.join(self.company.capabilities[:3])}. 
Our company holds certifications including {', '.join(self.company.company_certifications[:2])}, 
demonstrating our commitment to quality and security."""
        
        elif q_type == 'technical':
            return f"""Our technical approach leverages proven methodologies and modern technologies. 
We have successfully delivered similar projects for federal agencies including {', '.join([pp['client'] for pp in self.company.past_performance[:2]])}. 
Our core capabilities include {', '.join(self.company.capabilities[:3])}."""
        
        elif q_type == 'past_performance':
            refs = self.company.past_performance[:2]
            return f"""We have extensive relevant experience including:\n\n""" + "\n\n".join([
                f"{pp['project']} - {pp['client']} ({pp['period']})\n"
                f"Contract Value: {pp['value']}\n"
                f"Description: {pp['description']}\n"
                f"Reference: {pp['poc']}"
                for pp in refs
            ])
        
        elif q_type == 'personnel':
            return f"""Our proposed team includes experienced professionals:\n\n""" + "\n\n".join([
                f"{kp['name']}, {kp['title']}\n"
                f"Clearance: {kp['clearance']}\n"
                f"Experience: {kp['experience']}"
                for kp in self.company.key_personnel
            ])
        
        else:
            return """[Response to be completed based on specific question requirements. 
Please review and customize this section based on your company's specific qualifications and experience.]"""
    
    def generate_rfi_document(self, opportunity: Dict, questions: List[Dict], responses: List[str]) -> str:
        """Generate Word document with RFI responses"""
        
        if not DOCX_AVAILABLE:
            print("⚠️  python-docx not available, generating text file instead")
            return self.generate_text_document(opportunity, questions, responses)
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('Request for Information (RFI) Response', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Opportunity information
        doc.add_heading('Opportunity Information', 1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Opportunity Title:', opportunity.get('title', 'Unknown')),
            ('Agency:', opportunity.get('agency', 'Unknown')),
            ('Notice ID:', opportunity.get('notice_id', 'N/A')),
            ('Response Date:', datetime.now().strftime('%B %d, %Y'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            info_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Company information
        doc.add_heading('Respondent Information', 1)
        
        company_table = doc.add_table(rows=5, cols=2)
        company_table.style = 'Light Grid Accent 1'
        
        company_data = [
            ('Company Name:', self.company.company_name),
            ('DUNS Number:', self.company.duns),
            ('CAGE Code:', self.company.cage_code),
            ('Address:', self.company.address),
            ('Socioeconomic Status:', self.company.socioeconomic_status)
        ]
        
        for i, (label, value) in enumerate(company_data):
            company_table.rows[i].cells[0].text = label
            company_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            company_table.rows[i].cells[1].text = value
        
        doc.add_page_break()
        
        # Responses
        doc.add_heading('RFI Responses', 1)
        
        for question, response in zip(questions, responses):
            # Question
            q_heading = doc.add_heading(f"Question {question['number']}", 2)
            q_para = doc.add_paragraph(question['question'])
            q_para.runs[0].italic = True
            
            # Response
            doc.add_heading('Response:', 3)
            doc.add_paragraph(response)
            doc.add_paragraph()
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"rfi_response_{timestamp}.docx"
        doc.save(filename)
        
        return filename
    
    def generate_text_document(self, opportunity: Dict, questions: List[Dict], responses: List[str]) -> str:
        """Generate text file as fallback"""
        
        lines = []
        lines.append("="*70)
        lines.append("REQUEST FOR INFORMATION (RFI) RESPONSE")
        lines.append("="*70)
        lines.append("")
        
        lines.append("OPPORTUNITY INFORMATION")
        lines.append("-"*70)
        lines.append(f"Title: {opportunity.get('title', 'Unknown')}")
        lines.append(f"Agency: {opportunity.get('agency', 'Unknown')}")
        lines.append(f"Notice ID: {opportunity.get('notice_id', 'N/A')}")
        lines.append(f"Response Date: {datetime.now().strftime('%B %d, %Y')}")
        lines.append("")
        
        lines.append("RESPONDENT INFORMATION")
        lines.append("-"*70)
        lines.append(f"Company: {self.company.company_name}")
        lines.append(f"DUNS: {self.company.duns}")
        lines.append(f"CAGE Code: {self.company.cage_code}")
        lines.append(f"Address: {self.company.address}")
        lines.append(f"Status: {self.company.socioeconomic_status}")
        lines.append("")
        
        lines.append("="*70)
        lines.append("RFI RESPONSES")
        lines.append("="*70)
        lines.append("")
        
        for question, response in zip(questions, responses):
            lines.append(f"QUESTION {question['number']}")
            lines.append("-"*70)
            lines.append(question['question'])
            lines.append("")
            lines.append("RESPONSE:")
            lines.append(response)
            lines.append("")
            lines.append("")
        
        # Save text file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"rfi_response_{timestamp}.txt"
        
        with open(filename, 'w') as f:
            f.write('\n'.join(lines))
        
        return filename
    
    def generate_rfi_response(self, opportunity: Dict) -> str:
        """Complete RFI response generation"""
        
        print(f"\n🎯 Generating RFI Response")
        print("="*70)
        print(f"Opportunity: {opportunity.get('title', 'Unknown')}")
        print(f"Agency: {opportunity.get('agency', 'Unknown')}")
        print()
        
        # Extract questions
        print("📋 Extracting RFI questions...")
        questions = self.extract_rfi_questions(opportunity)
        print(f"   Found {len(questions)} questions")
        print()
        
        # Generate responses
        print("✍️  Generating responses...")
        responses = []
        
        for i, question in enumerate(questions, 1):
            print(f"   Question {i}/{len(questions)}: {question['question'][:50]}...")
            
            if self.claude_available:
                response = self.generate_response_with_claude(question, opportunity)
            else:
                response = self.generate_fallback_response(question)
            
            responses.append(response)
        
        print()
        
        # Generate document
        print("📄 Creating RFI document...")
        filename = self.generate_rfi_document(opportunity, questions, responses)
        
        print(f"✓ RFI response saved to: {filename}")
        print()
        
        return filename


def main():
    """Test the RFI generator"""
    
    print("\n📝 RFI RESPONSE GENERATOR")
    print("="*70)
    
    # Test opportunity
    test_opportunity = {
        'title': 'Cloud Migration and Cybersecurity Services',
        'agency': 'Department of Defense',
        'notice_id': 'TEST-2026-001',
        'description': '''
        Request for Information
        
        The Department of Defense is seeking information from qualified contractors
        for cloud migration and cybersecurity services.
        
        Please respond to the following questions:
        
        1. Describe your company's experience with federal cloud migration projects,
        including specific examples from the past 3 years.
        
        2. What is your technical approach to securing cloud environments in compliance
        with DoD security requirements?
        
        3. Provide information about your key personnel who would support this effort,
        including their clearance levels and relevant certifications.
        
        4. What is your company's socioeconomic status and DUNS number?
        '''
    }
    
    generator = RFIResponseGenerator()
    filename = generator.generate_rfi_response(test_opportunity)
    
    print(f"\n🎉 Complete! Open the file to review:")
    print(f"   {filename}")
    print()


if __name__ == '__main__':
    main()
